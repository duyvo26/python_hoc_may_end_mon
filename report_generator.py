import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    def __init__(self, session_dir):
        self.session_dir = session_dir
        self.report_md_path = os.path.join(session_dir, "Full_Report.md")
        self.report_docx_path = os.path.join(session_dir, "Full_Report.docx")

    def generate(self, dataset_name, original_info, preprocess_info, k_details, voting_history, best_k_msg, metrics_df, profile_km, profile_h):
        """Tạo báo cáo .md và .docx"""
        
        # 1. TẠO NỘI DUNG MARKDOWN
        try:
            md_content = f"""# BÁO CÁO PHÂN TÍCH PHÂN CỤM DỮ LIỆU
**Tên bộ dữ liệu:** {dataset_name}  
**Ngày thực hiện:** {os.path.basename(self.session_dir)}

## 1. Thông tin dữ liệu và Tiền xử lý
- **Dữ liệu gốc:** {original_info}
- **Kết quả tiền xử lý:** {preprocess_info}

## 2. Kết quả phân tích số cụm tối ưu (K)
### Đồng thuận:
{best_k_msg}

### Chỉ số đánh giá (Lần cuối):
{k_details.to_markdown(index=False) if hasattr(k_details, 'to_markdown') else k_details}

### Nhật ký biểu quyết (N Trials):
{voting_history.to_markdown(index=False) if hasattr(voting_history, 'to_markdown') else voting_history}

![Phân tích KMeans](1_Analysis_KMeans.png)
![Phân tích Hierarchical](1_Analysis_Hierarchical.png)

## 3. Kết quả huấn luyện mô hình
### So sánh hiệu năng giữa 2 thuật toán:
{metrics_df.to_markdown(index=False) if hasattr(metrics_df, 'to_markdown') else metrics_df}

![PCA 2D KMeans](2_PCA_KMeans_2D.png)
![PCA 2D Hierarchical](2_PCA_Hierarchical_2D.png)

## 4. Đặc trưng của các cụm (Profiling)
### Đặc trưng K-Means:
{profile_km.to_markdown(index=False) if hasattr(profile_km, 'to_markdown') else profile_km}

### Đặc trưng Hierarchical:
{profile_h.to_markdown(index=False) if hasattr(profile_h, 'to_markdown') else profile_h}
"""
            with open(self.report_md_path, "w", encoding="utf-8") as f:
                f.write(md_content)
        except Exception as e:
            print(f"Error generating MD: {e}")

        # 2. CHUYỂN ĐỔI SANG DOCX
        try:
            self._create_docx(dataset_name, original_info, preprocess_info, k_details, voting_history, best_k_msg, metrics_df, profile_km, profile_h)
        except Exception as e:
            print(f"Error generating DOCX: {e}")
        
        return self.report_md_path, self.report_docx_path

    def _create_docx(self, dataset_name, original_info, preprocess_info, k_details, voting_history, best_k_msg, metrics_df, profile_km, profile_h):
        doc = Document()
        
        # Tiêu đề chính
        title = doc.add_heading('BÁO CÁO PHÂN TÍCH PHÂN CỤM', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Tên bộ dữ liệu: {dataset_name}")
        doc.add_paragraph(f"Phiên làm việc: {os.path.basename(self.session_dir)}")
        
        # Section 1: Data Info
        doc.add_heading('1. Thông tin dữ liệu và Tiền xử lý', level=1)
        doc.add_paragraph(f"Dữ liệu gốc: {original_info}")
        doc.add_paragraph(f"Kết quả tiền xử lý: {preprocess_info}")

        # Section 2: K Analysis
        doc.add_heading('2. Phân tích số cụm tối ưu (K)', level=1)
        doc.add_paragraph(f"Kết quả đồng thuận: {best_k_msg}")
        
        def add_table(df, table_title):
            doc.add_heading(table_title, level=2)
            if df.empty:
                doc.add_paragraph("Không có dữ liệu.")
                return
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        add_table(k_details, "Chỉ số đánh giá (Lần cuối)")
        add_table(voting_history, "Nhật ký biểu quyết (N Trials)")
        
        # Chèn ảnh phân tích K
        for img in ["1_Analysis_KMeans.png", "1_Analysis_Hierarchical.png"]:
            img_path = os.path.join(self.session_dir, img)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
        
        # Section 3
        doc.add_heading('3. Kết quả huấn luyện và PCA', level=1)
        add_table(metrics_df, "Bảng so sánh Metrics hiệu năng")
        # Chèn các ảnh PCA
        for img in ["2_PCA_KMeans_2D.png", "2_PCA_Hierarchical_2D.png"]:
            img_path = os.path.join(self.session_dir, img)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
        
        # Section 4
        doc.add_heading('4. Đặc trưng các cụm (Profiling)', level=1)
        add_table(profile_km, "Bảng Profiling K-Means")
        add_table(profile_h, "Bảng Profiling Hierarchical")
        
        doc.save(self.report_docx_path)
